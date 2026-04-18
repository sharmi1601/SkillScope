"""
resume_parser.py — Turn a user's resume (PDF or DOCX) into a SkillScope profile JSON.

Flow:
  1. Extract raw text from the file (PyMuPDF for PDF, python-docx for DOCX).
  2. Feed text to Gemini with a strict extraction prompt that:
       - Pulls out concrete, teachable skills (same rules as the JD extractor)
       - Assigns a confidence level (strong / basic) based on how the skill
         appears in the resume (multiple contexts → strong, brief mention → basic)
       - Captures experiences, projects, education
       - Maps skill names to the same canonical forms the aggregator uses
  3. Merges the LLM output with user-supplied metadata (target_role, hours_per_week,
     deadline_weeks, learning_style, name) to produce a profile compatible with
     example_profile.json.

The output profile is immediately usable by:
    python -m src.score_cli --role <role> --profile data/profile_<name>.json --plan
"""

from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any

from dotenv import load_dotenv

load_dotenv()

# Same provider + model defaults as the JD extractor so the whole system is
# consistent. Users who want a smarter model for resume parsing can set
# GEMINI_MODEL_RESUME independently.
_PROVIDER = os.getenv("LLM_PROVIDER", "gemini").lower()
_GEMINI_MODEL = os.getenv(
    "GEMINI_MODEL_RESUME",
    os.getenv("GEMINI_MODEL_EXTRACTION", "gemini-2.5-flash-lite"),
)
_GROQ_MODEL = os.getenv("GROQ_MODEL_EXTRACTION", "llama-3.3-70b-versatile")


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


def extract_text(path: str | Path) -> str:
    """Dispatch by extension. Returns plain text concatenated across pages/paragraphs."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Resume not found: {p}")

    ext = p.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(p)
    if ext in (".docx", ".doc"):
        return _extract_docx(p)
    if ext in (".txt", ".md"):
        return p.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(
        f"Unsupported resume format: {ext}. Use .pdf, .docx, .txt, or .md."
    )


def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF using PyMuPDF (fitz). Fast, accurate, no OCR."""
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise RuntimeError(
            "pymupdf is required for PDF resumes. Install with:\n"
            "    pip install pymupdf"
        ) from e

    parts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts).strip()


def _extract_docx(path: Path) -> str:
    """Extract paragraph text from a DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for DOCX resumes. Install with:\n"
            "    pip install python-docx"
        ) from e

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables often hold skill lists on modern resume templates.
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    parts.append(txt)
    return "\n".join(parts).strip()


# ---------------------------------------------------------------------------
# Gemini prompt — mirrors the JD extractor's rigor
# ---------------------------------------------------------------------------


SYSTEM_PROMPT = """You are a precise technical recruiter. Your job is to parse a candidate's resume and extract a STRUCTURED profile of their concrete, teachable skills and background.

## What counts as a skill

A "skill" is a SPECIFIC, NAMEABLE, LEARNABLE capability:
- Programming languages: Python, SQL, Java, TypeScript, Go
- Libraries/frameworks: React, Django, PyTorch, pandas, Spring Boot
- Tools/platforms: Tableau, Airflow, Docker, Kubernetes, AWS, Snowflake, dbt
- Techniques: A/B Testing, Causal Inference, ETL, REST API design

## What does NOT count — NEVER extract these

Fields / soft skills / credentials:
  AI, Data Science, Software Engineering, Analytics, DevOps,
  teamwork, communication, leadership, passion, problem solving,
  Bachelor's degree, Master's degree, "5 years experience"

## Confidence rules — IMPORTANT

For each skill, assign one of two confidence levels:
- "strong"  → The skill appears in MULTIPLE contexts: a job title, a project,
             a responsibilities bullet, AND/OR a skills list. Or the resume
             explicitly says "expert", "proficient", "lead", "primary".
- "basic"   → Mentioned once, listed in a skills section only, or phrased as
             "familiar with", "exposure to", "introduction to".

When in doubt, prefer "basic" — overclaiming confidence leads to bad gap reports.

## Normalize skill names

- "Postgres" | "postgresql" → "PostgreSQL"
- "JS" | "Javascript"       → "JavaScript"
- "ML"                      → "Machine Learning"
- "NLP"                     → "Natural Language Processing"
- Lowercase for Python package names: "pandas", "numpy", "dbt", "scikit-learn"
- Title Case for technique names: "A/B Testing", "Causal Inference"

## Output schema — STRICT JSON, no extra keys, no prose

{
  "name": "string — candidate's full name, or 'user' if not found",
  "skills": [
    {"name": "string", "confidence": "strong | basic"}
  ],
  "experiences": [
    {"title": "string", "company": "string", "years": number}
  ],
  "projects": [
    {"name": "string", "tech": ["skill1", "skill2"]}
  ],
  "education": [
    {"degree": "string", "field": "string"}
  ]
}

If any section is absent from the resume, return an empty list for it.
Return JSON only. No prose, no markdown fences.
"""


USER_TEMPLATE = """Resume text:
---
{resume_text}
---

Extract the structured profile. Respond with JSON only."""


# ---------------------------------------------------------------------------
# LLM call — same provider-dispatch pattern as extractor.py / scheduler.py
# ---------------------------------------------------------------------------


def _call_gemini(resume_text: str) -> dict[str, Any]:
    from google import genai
    from google.genai import types

    key = os.getenv("GEMINI_API_KEY")
    if not key:
        raise RuntimeError("GEMINI_API_KEY not set in .env")
    client = genai.Client(api_key=key)
    # Truncate absurdly long resumes (some Europass CVs go 20+ pages).
    trimmed = resume_text[:16000]
    resp = client.models.generate_content(
        model=_GEMINI_MODEL,
        contents=USER_TEMPLATE.format(resume_text=trimmed),
        config=types.GenerateContentConfig(
            system_instruction=SYSTEM_PROMPT,
            temperature=0.1,
            response_mime_type="application/json",
        ),
    )
    return json.loads(resp.text or "{}")


def _call_groq(resume_text: str) -> dict[str, Any]:
    from groq import Groq

    key = os.getenv("GROQ_API_KEY")
    if not key:
        raise RuntimeError("GROQ_API_KEY not set in .env")
    client = Groq(api_key=key)
    trimmed = resume_text[:16000]
    resp = client.chat.completions.create(
        model=_GROQ_MODEL,
        response_format={"type": "json_object"},
        temperature=0.1,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": USER_TEMPLATE.format(resume_text=trimmed)},
        ],
    )
    return json.loads(resp.choices[0].message.content or "{}")


def _call_llm(resume_text: str) -> dict[str, Any]:
    if _PROVIDER == "gemini":
        return _call_gemini(resume_text)
    if _PROVIDER == "groq":
        return _call_groq(resume_text)
    raise RuntimeError(f"Unknown LLM_PROVIDER={_PROVIDER!r}. Use 'gemini' or 'groq'.")


# ---------------------------------------------------------------------------
# Top-level parse function
# ---------------------------------------------------------------------------


def parse_resume(
    path: str | Path,
    target_role: str,
    name: str | None = None,
    hours_per_week: int = 10,
    deadline_weeks: int = 12,
    learning_style: str = "any",
) -> dict[str, Any]:
    """
    Parse a resume file into a SkillScope profile.

    Parameters not available in the resume (hours_per_week, deadline_weeks,
    learning_style, target_role) are supplied by the caller — these come from
    a separate form in the UI, not the resume itself.

    Returns a profile dict matching the schema of data/example_profile.json.
    """
    text = extract_text(path)
    if not text:
        raise RuntimeError(
            f"Extracted no text from {path}. Is it a scanned image PDF? "
            "(this parser does not OCR)."
        )

    extracted = _call_llm(text)

    # Merge LLM-extracted fields with caller-supplied metadata.
    profile = {
        "name": name or extracted.get("name") or "user",
        "target_role": target_role,
        "hours_per_week": hours_per_week,
        "deadline_weeks": deadline_weeks,
        "learning_style": learning_style,
        "skills": extracted.get("skills", []),
        "experiences": extracted.get("experiences", []),
        "projects": extracted.get("projects", []),
        "education": extracted.get("education", []),
        # Keep a copy of the first ~500 chars of the raw resume for debugging.
        "_source_excerpt": text[:500],
    }
    return profile


if __name__ == "__main__":
    # Smoke test: parse a plain-text "resume" string through the Gemini path.
    fake = """
    Sharmi Desiboyina
    Software Engineer, 2 years experience

    SKILLS
    - Python (expert, daily for 2 years)
    - SQL (intermediate, used at 2 jobs)
    - pandas, numpy
    - matplotlib, seaborn
    - Git, Docker
    - Familiar with: PostgreSQL, Airflow

    EXPERIENCE
    Software Engineer, Acme Corp (2023-2025)
    - Built ETL pipelines in Python + Airflow
    - Wrote SQL reporting queries on PostgreSQL

    PROJECTS
    - Portfolio analytics dashboard (Python, pandas, matplotlib)
    - Personal finance tracker (Python, SQLite)

    EDUCATION
    B.Tech, Computer Science, 2023
    """

    import tempfile

    with tempfile.NamedTemporaryFile("w", suffix=".txt", delete=False) as f:
        f.write(fake)
        path = f.name

    profile = parse_resume(
        path,
        target_role="data_analyst",
        name="Sharmi",
        hours_per_week=10,
        deadline_weeks=12,
        learning_style="video",
    )
    print(json.dumps(profile, indent=2))
